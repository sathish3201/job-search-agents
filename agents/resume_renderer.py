"""Renders a structured TailoredDraft into real PDF/DOCX bytes.

Single responsibility: structured draft in, bytes out — no knowledge of
chat sessions, scoring, or agents. Mirrors document_parser.py's
"bytes-in, text-out" mandate, in reverse.

Scope decision (see plan): this renders ONE clean, consistent template for
every user, not a reproduction of the user's original PDF's specific
design — there's no way to reverse-engineer an arbitrary uploaded PDF's
fonts/layout, and that's a different, much harder problem. The left pane
of the tailoring dialog shows the user's real original file as ground
truth; this produces the right pane's tailored preview/export.
"""
from __future__ import annotations

import io

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

from models import TailoredDraft

_HEADING_TYPES = {"experience_heading", "education", "other"}


def render_pdf(draft: TailoredDraft, candidate_name: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("NameStyle", parent=styles["Title"], fontSize=18, spaceAfter=4)
    heading_style = ParagraphStyle(
        "SectionHeading", parent=styles["Heading2"], spaceBefore=10, spaceAfter=4
    )
    body_style = ParagraphStyle("Body", parent=styles["Normal"], spaceAfter=4)

    story = [Paragraph(_escape(candidate_name), name_style), Spacer(1, 6)]

    bullets: list[str] = []

    def flush_bullets():
        if bullets:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_escape(b), body_style)) for b in bullets],
                    bulletType="bullet",
                )
            )
            bullets.clear()

    for section in sorted(draft.sections, key=lambda s: s.order):
        if section.section_type in ("headline", "summary"):
            flush_bullets()
            story.append(Paragraph(_escape(section.text), body_style))
        elif section.section_type == "experience_bullet":
            bullets.append(section.text)
        elif section.section_type in _HEADING_TYPES:
            flush_bullets()
            story.append(Paragraph(_escape(section.text), heading_style))
        elif section.section_type == "skill":
            bullets.append(section.text)
        else:
            flush_bullets()
            story.append(Paragraph(_escape(section.text), body_style))
    flush_bullets()

    doc.build(story)
    return buf.getvalue()


def render_docx(draft: TailoredDraft, candidate_name: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(candidate_name, level=0)

    for section in sorted(draft.sections, key=lambda s: s.order):
        if section.section_type in ("headline", "summary"):
            doc.add_paragraph(section.text)
        elif section.section_type in ("experience_bullet", "skill"):
            doc.add_paragraph(section.text, style="List Bullet")
        elif section.section_type in _HEADING_TYPES:
            doc.add_heading(section.text, level=2)
        else:
            doc.add_paragraph(section.text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _escape(text: str) -> str:
    # reportlab's Paragraph interprets a minimal XML-like markup — user/LLM
    # text containing literal &, <, > must be escaped or it silently breaks
    # rendering (or is misinterpreted as markup).
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
